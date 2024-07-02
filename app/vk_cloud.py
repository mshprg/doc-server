import asyncio
import os
import time
import requests
import docx
from pdf2image import convert_from_bytes
import io

lock = asyncio.Lock()
access_token = "6QNmzdpQrmR3wavR8WphKDjiDq8Uaf2FBUvRkWbBdCLNvRNZw"
refresh_token = os.environ.get("VK_REFRESH_TOKEN")
token_expired = 0


def get_access_token(client_id, refresh_token):
    url = "https://mcs.mail.ru/auth/oauth/v1/token"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "client_id": client_id,
        "refresh_token": refresh_token,
        "grant_type": "refresh_token"
    }
    response = requests.post(url, headers=headers, json=data)
    return response.json()


async def send_image(file):
    global token_expired, access_token
    if time.time() >= token_expired:
        async with lock:
            refresh_tok()
    url = f'https://smarty.mail.ru/api/v1/text/recognize?oauth_token={access_token}&oauth_provider=mcs'

    headers = {
        'accept': 'application/json',
    }

    files = {
        'file': (f'MyImage.jpeg', file, 'image/jpeg'),
    }

    data = {
        'meta': '{"images":[{"name":"file"}]}'
    }

    response = requests.post(url, headers=headers, files=files, data=data)
    if response.status_code == 200:
        return response.json()['body']['objects'][0]['text']
    else:
        return None


async def get_text(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text), len(doc.paragraphs)


async def pdf_to_img(file):
    images = convert_from_bytes(pdf_file=file.read())
    output_str = ''
    for image in images:
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        output_str += await send_image(img_byte_arr)
    return output_str, len(images)


def refresh_tok():
    global access_token, token_expired
    access_token = get_access_token(client_id, refresh_token)['access_token']
    token_expired = time.time() + 3600
    return


client_id = os.environ.get("VK_CLIENT_ID")
client_secret = os.environ.get("VK_CLIENT_SECRET")
